from __future__ import annotations

from io import BytesIO
from datetime import datetime
from typing import Any
import html
import base64

import pandas as pd

from services.ratio_engine import format_ratio


def build_board_sections(state: dict, ratios: pd.DataFrame, inputs: dict) -> list[dict]:
    profile = state.get("company_profile", {}) or {}
    company = profile.get("Company Name", "Company")
    period = profile.get("Report Period") or profile.get("Financial Year", "Current period")
    commentary = state.get("board_ai_narrative") or state.get("ai_commentary") or "AI commentary has not yet been generated."
    risks = inputs.get("Top risks", "Not provided")
    priorities = inputs.get("Strategic priorities", "Not provided")
    decisions = inputs.get("Decisions required", "No board decisions recorded.")
    outlook = inputs.get("Management outlook", "Management outlook has not been entered.")
    people = state.get("board_people_data", {}) or {}
    ratio_lines = []
    currency = profile.get("Currency", "AUD")
    for _, row in ratios.iterrows():
        ratio_lines.append(f"{row['Ratio']}: {format_ratio(row['Value'], row['Unit'], currency)} ({row['Status']})")
    return [
        {"title": "Executive summary", "body": commentary},
        {"title": "Performance and outlook", "body": outlook},
        {"title": "Financial ratios and working capital", "body": "\n".join(ratio_lines)},
        {"title": "People and operating capacity", "body": (
            f"Total employees: {int(people.get('Total employees', 0) or 0)}; technical: {int(people.get('Technical staff', 0) or 0)}; "
            f"management: {int(people.get('Management staff', 0) or 0)}; apprentices: {int(people.get('Apprentice staff', 0) or 0)}. "
            f"Operational commentary: {inputs.get('People commentary', 'Not provided')}"
        )},
        {"title": "Strategic priorities", "body": priorities},
        {"title": "Principal risks and mitigations", "body": risks},
        {"title": "Board decisions and approvals", "body": decisions},
        {"title": "Governance confirmation", "body": inputs.get("Governance commentary", "No additional governance matters recorded.")},
    ]


def create_board_report_docx(state: dict, ratios: pd.DataFrame, inputs: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    profile = state.get("company_profile", {}) or {}
    company = profile.get("Company Name", "Company")
    period = profile.get("Report Period") or profile.get("Financial Year", "Current period")
    currency = profile.get("Currency", "AUD")
    sections = build_board_sections(state, ratios, inputs)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65)
    sec.left_margin = Inches(.72); sec.right_margin = Inches(.72)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(9.5)
    for style_name, size, colour in [("Title", 25, "17345B"), ("Heading 1", 16, "17345B"), ("Heading 2", 12, "2F5597")]:
        style = styles[style_name]
        style.font.name = "Arial"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(colour)

    logo_bytes = state.get("company_logo_bytes")
    if logo_bytes:
        try:
            logo_stream = BytesIO(logo_bytes)
            p_logo = doc.add_paragraph(); p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_stream, width=Inches(1.65))
        except Exception:
            pass
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BOARD & MANAGEMENT REPORT"); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(76, 116, 196)
    title = doc.add_paragraph(company, style="Title"); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"{period}  |  Prepared {datetime.now():%d %B %Y}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    # Board dashboard table
    key_ratios = ratios[ratios["Ratio"].isin(["Gross Margin", "Net Profit Margin", "Current Ratio", "DSO", "Cash Conversion Cycle", "Debt to Equity", "Revenue per Employee"])]
    table = doc.add_table(rows=1, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
    headers = ["Key measure", "Result", "Assessment"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = text; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs: run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "2F5597"); cell._tc.get_or_add_tcPr().append(shd)
    for _, row in key_ratios.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Ratio"])
        cells[1].text = format_ratio(row["Value"], row["Unit"], currency)
        cells[2].text = str(row["Status"])
    doc.add_paragraph()

    for section in sections:
        doc.add_heading(section["title"], level=1)
        for line in str(section["body"]).splitlines():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(5)

    doc.add_heading("Appendix A - Full ratio register", level=1)
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(["Ratio", "Category", "Result", "Status"]):
        table.rows[0].cells[i].text = text
        for run in table.rows[0].cells[i].paragraphs[0].runs: run.bold = True
    for _, row in ratios.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Ratio"]); cells[1].text = str(row["Category"])
        cells[2].text = format_ratio(row["Value"], row["Unit"], currency); cells[3].text = str(row["Status"])

    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Confidential - AI CFO Copilot board pack").font.size = Pt(8)
    out = BytesIO(); doc.save(out); return out.getvalue()


def create_board_report_html(state: dict, ratios: pd.DataFrame, inputs: dict) -> bytes:
    profile = state.get("company_profile", {}) or {}
    company = html.escape(str(profile.get("Company Name", "Company")))
    period = html.escape(str(profile.get("Report Period") or profile.get("Financial Year", "Current period")))
    currency = profile.get("Currency", "AUD")
    sections = build_board_sections(state, ratios, inputs)
    logo_bytes = state.get("company_logo_bytes")
    logo_html = ""
    if logo_bytes:
        try:
            encoded = base64.b64encode(logo_bytes).decode("ascii")
            name = str(state.get("company_logo_name") or "").lower()
            mime = "image/jpeg" if name.endswith((".jpg", ".jpeg")) else "image/png"
            logo_html = f"<img src='data:{mime};base64,{encoded}' alt='Company logo' style='max-height:82px;max-width:240px;margin-bottom:18px;background:white;padding:8px;border-radius:12px'>"
        except Exception:
            logo_html = ""
    ratio_rows = "".join(
        f"<tr><td>{html.escape(str(r['Ratio']))}</td><td>{html.escape(str(r['Category']))}</td><td>{html.escape(format_ratio(r['Value'], r['Unit'], currency))}</td><td>{html.escape(str(r['Status']))}</td></tr>"
        for _, r in ratios.iterrows()
    )
    section_html = "".join(f"<section><h2>{html.escape(s['title'])}</h2><p>{html.escape(str(s['body'])).replace(chr(10), '<br>')}</p></section>" for s in sections)
    document = f"""<!doctype html><html><head><meta charset='utf-8'><title>{company} Board Report</title><style>
    body{{font-family:Arial,sans-serif;color:#172033;max-width:1000px;margin:40px auto;line-height:1.5}}header{{padding:32px;background:linear-gradient(135deg,#17345b,#4a6fc1);color:white;border-radius:18px}}h1{{margin:0}}h2{{color:#17345b;border-bottom:2px solid #dbe5f5;padding-bottom:7px}}section{{margin:30px 0}}table{{width:100%;border-collapse:collapse}}th,td{{border:1px solid #d7deea;padding:9px;text-align:left}}th{{background:#eef3fb}}.meta{{opacity:.85}}</style></head><body>
    <header>{logo_html}<div>BOARD & MANAGEMENT REPORT</div><h1>{company}</h1><div class='meta'>{period} | Prepared {datetime.now():%d %B %Y}</div></header>{section_html}
    <section><h2>Full ratio register</h2><table><thead><tr><th>Ratio</th><th>Category</th><th>Result</th><th>Status</th></tr></thead><tbody>{ratio_rows}</tbody></table></section>
    <p><small>Confidential - generated by AI CFO Copilot. Management must validate source data, assumptions and AI-generated narrative before circulation.</small></p></body></html>"""
    return document.encode("utf-8")
